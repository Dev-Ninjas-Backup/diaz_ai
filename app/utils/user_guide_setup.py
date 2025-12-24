from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
from uuid import uuid4

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams

from app.services.chatbot.llms.open_ai_llm import OpenaiLLM
from app.config import configs
from app.utils.logger import get_logger

from docx import Document as DocxDocument
import io

router = APIRouter()
logger = get_logger(__name__)


@router.post("/rag/upload-docx")
async def upload_docx_to_qdrant(
    file: UploadFile = File(...),
    source: str = Form(None)
):
    try:
        collection_name = "florida_yacht_sales"
        logger.info(f"Uploading DOCX document to collection: {collection_name}")

        if not file.filename.endswith(".docx"):
            raise HTTPException(status_code=400, detail="Only .docx files are supported")

        # -----------------------------
        # Read DOCX file
        # -----------------------------
        file_content = await file.read()
        docx_file = DocxDocument(io.BytesIO(file_content))

        text_content = "\n".join([p.text for p in docx_file.paragraphs if p.text.strip()])
        if not text_content:
            raise HTTPException(status_code=400, detail="DOCX file is empty")

        # -----------------------------
        # Chunking
        # -----------------------------
        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        chunks = splitter.split_text(text_content)

        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": source or file.filename,
                    "doc_id": str(uuid4())
                }
            )
            for chunk in chunks
        ]

        # -----------------------------
        # Qdrant setup
        # -----------------------------
        client = QdrantClient(url=configs["vector"]["url"])
        embeddings = OpenaiLLM().get_embeddings()

        # create collection if not exists
        existing_collections = [c.name for c in client.get_collections().collections]
        if collection_name not in existing_collections:
            client.create_collection(
                collection_name=collection_name,
                vectors_config=VectorParams(
                    size=3072,  # OpenAI text-embedding-3-large dimension
                    distance=Distance.COSINE
                )
            )
            logger.info(f"Collection '{collection_name}' created successfully")
        else:
            logger.info(f"Collection '{collection_name}' already exists")

        vectorstore = QdrantVectorStore(
            client=client,
            collection_name=collection_name,
            embedding=embeddings
        )

        # -----------------------------
        # Store documents
        # -----------------------------
        vectorstore.add_documents(documents)

        logger.info(f"Uploaded {len(documents)} chunks from DOCX file: {file.filename}")

        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "collection": collection_name,
                "chunks_uploaded": len(documents),
                "file": file.filename
            }
        )

    except Exception as e:
        logger.error(f"Failed to upload DOCX docs: {e}")
        raise HTTPException(status_code=500, detail=str(e))